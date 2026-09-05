"""
Resume Parser Service
Supports PDF (pdfplumber) and DOCX (python-docx) formats.
"""

from __future__ import annotations

import io
import re
import unicodedata
from pathlib import Path


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------
def parse_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF byte string using pdfplumber (with pypdf fallback)."""
    import pdfplumber  # lazy import so the module is optional at import time

    pages: list[str] = []
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if not text or not text.strip():
                    text = page.extract_text()
                if text and text.strip():
                    pages.append(text.strip())
    except Exception:
        pass

    # Fallback to pypdf if pdfplumber extracted nothing or encountered an error
    if not pages:
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            for page in reader.pages:
                t = page.extract_text()
                if t and t.strip():
                    pages.append(t.strip())
        except Exception:
            pass

    return "\n\n".join(pages)


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------
def parse_docx(file_bytes: bytes) -> str:
    """Extract plain text from a DOCX byte string using python-docx."""
    from docx import Document  # lazy import

    doc = Document(io.BytesIO(file_bytes))
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    paragraphs.append(text)

    return "\n".join(paragraphs)


# ---------------------------------------------------------------------------
# Text normalisation
# ---------------------------------------------------------------------------
def parse_text(text: str) -> str:
    """Clean and normalise extracted resume text."""
    if not text:
        return ""

    # Normalise unicode characters (e.g. smart quotes, em-dashes)
    text = unicodedata.normalize("NFKD", text)

    # Replace non-breaking spaces and other whitespace variants
    text = text.replace("\xa0", " ").replace("\t", " ")

    # Collapse runs of blank lines (keep at most 2)
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Collapse runs of spaces inside a line
    lines = [re.sub(r" {2,}", " ", line) for line in text.splitlines()]

    # Strip trailing/leading whitespace per line
    lines = [line.strip() for line in lines]

    return "\n".join(lines).strip()


# ---------------------------------------------------------------------------
# Extension helper
# ---------------------------------------------------------------------------
def get_file_extension(filename: str) -> str:
    """Return the lower-case file extension without the leading dot."""
    return Path(filename).suffix.lstrip(".").lower()


# ---------------------------------------------------------------------------
# Public entry-point
# ---------------------------------------------------------------------------
def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Detect the file type from *filename*, route to the appropriate parser,
    and return normalised plain text.

    Raises ValueError for unsupported file types.
    """
    ext = get_file_extension(filename)

    if ext == "pdf":
        raw_text = parse_pdf(file_bytes)
    elif ext in ("docx", "doc"):
        raw_text = parse_docx(file_bytes)
    elif ext == "txt":
        raw_text = file_bytes.decode("utf-8", errors="replace")
    else:
        raise ValueError(
            f"Unsupported file type: '.{ext}'. "
            "Please upload a PDF, DOCX, or TXT file."
        )

    return parse_text(raw_text)
